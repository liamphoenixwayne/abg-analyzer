#!/usr/bin/env python3
"""GitHub Light Word 引擎 v6 — gen: Markdown→docx；restyle: 格式化现有 docx"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

FONT_LATIN = 'Segoe UI'
FONT_CJK = '微软雅黑'
FONT_MONO = 'Consolas'

INK     = RGBColor(0x1f, 0x23, 0x28)
GREY    = RGBColor(0x59, 0x63, 0x6e)
ACCENT  = RGBColor(0x09, 0x69, 0xda)
SUCCESS = RGBColor(0x1a, 0x7f, 0x37)
WARN    = RGBColor(0x9a, 0x67, 0x00)
DANGER  = RGBColor(0xd1, 0x24, 0x2f)
PURPLE  = RGBColor(0x82, 0x50, 0xdf)

BG, CODE_BG, BORDER = 'F6F8FA', 'F0F1F2', 'D1D9E0'
BODY_PT, CODE_PT, TABLE_PT, AFTER_TW = 12, 10, 10.5, '240'
FOOTER_TEXT = 'Designed by 葱头爸爸 · liamphoenixwayne@outlook.com'

CALLOUTS = {
    'NOTE':      (ACCENT,  '0969DA', 'Note',      'DDF4FF'),
    'TIP':       (SUCCESS, '1A7F37', 'Tip',       'DAFBE1'),
    'IMPORTANT': (PURPLE,  '8250DF', 'Important', 'FBEFFF'),
    'WARNING':   (WARN,    '9A6700', 'Warning',   'FFF8C5'),
    'CAUTION':   (DANGER,  'D1242F', 'Caution',   'FFEBE9'),
}
COLOR_MARKERS = [
    (re.compile(r'^\s*(⚠️?|【?(注意|警告)】?[:：]?|WARNING[:：]?)', re.I), 'WARNING'),
    (re.compile(r'^\s*(❗|【?(危险|禁忌|慎用)】?[:：]?|CAUTION[:：]?)', re.I), 'CAUTION'),
    (re.compile(r'^\s*(💡|【?(提示|技巧|建议)】?[:：]?|TIP[:：]?)', re.I), 'TIP'),
    (re.compile(r'^\s*(【?重要】?[:：]?|IMPORTANT[:：]?)', re.I), 'IMPORTANT'),
    (re.compile(r'^\s*(ℹ️?|【?(说明|备注)】?[:：]?|NOTE[:：]?)', re.I), 'NOTE'),
]
HEADING_SPECS = {1: (24.0, INK, True), 2: (18.0, INK, True), 3: (15.0, INK, False),
                 4: (12.0, INK, False), 5: (10.5, INK, False), 6: (10.0, GREY, False)}


def _shd(parent, fill):
    for old in parent.findall(qn('w:shd')):
        parent.remove(old)
    e = OxmlElement('w:shd')
    e.set(qn('w:val'), 'clear')
    e.set(qn('w:color'), 'auto')
    e.set(qn('w:fill'), fill)
    parent.append(e)


def _pbdr(ppr, edges, color, sz, space):
    for old in ppr.findall(qn('w:pBdr')):
        ppr.remove(old)
    pbdr = OxmlElement('w:pBdr')
    for edge in edges:
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), str(space))
        e.set(qn('w:color'), color)
        pbdr.append(e)
    ppr.append(pbdr)


def _rfonts(rpr, latin, cjk):
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    rf.set(qn('w:ascii'), latin)
    rf.set(qn('w:hAnsi'), latin)
    rf.set(qn('w:eastAsia'), cjk)


def _spacing(ppr, before='0', after=AFTER_TW, line: Optional[str] = '276'):
    for old in ppr.findall(qn('w:spacing')):
        ppr.remove(old)
    e = OxmlElement('w:spacing')
    e.set(qn('w:before'), before)
    e.set(qn('w:after'), after)
    if line:
        e.set(qn('w:line'), line)
        e.set(qn('w:lineRule'), 'auto')
    ppr.append(e)


def _ensure_style(doc, name, stype, base=None, builtin=False):
    try:
        return doc.styles[name]
    except KeyError:
        st = doc.styles.add_style(name, stype, builtin=builtin)
        if base:
            try:
                st.base_style = doc.styles[base]
            except KeyError:
                pass
        return st


def apply_github_styles(doc, cjk=FONT_CJK, line_mult=1.15, compact=False):
    line_tw = str(int(round(line_mult * 240)))
    after_tw = '120' if compact else AFTER_TW
    h_before = '240' if compact else '360'

    dd = doc.styles.element.find(qn('w:docDefaults'))
    if dd is not None:
        rprd = dd.find(qn('w:rPrDefault'))
        if rprd is not None:
            rpr = rprd.find(qn('w:rPr'))
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                rprd.append(rpr)
            _rfonts(rpr, FONT_LATIN, cjk)

    n = doc.styles['Normal']
    n.font.name = FONT_LATIN
    n.font.size = Pt(BODY_PT)
    n.font.color.rgb = INK
    _rfonts(n.element.get_or_add_rPr(), FONT_LATIN, cjk)
    _spacing(n.element.get_or_add_pPr(), after=after_tw, line=line_tw)

    def _hstyle(st, size, color, rule):
        st.font.name = FONT_LATIN
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.italic = False
        st.font.color.rgb = color
        _rfonts(st.element.get_or_add_rPr(), FONT_LATIN, cjk)
        ppr = st.element.get_or_add_pPr()
        _spacing(ppr, before=h_before, after=AFTER_TW, line='300')
        for old in ppr.findall(qn('w:pBdr')):
            ppr.remove(old)
        if rule:
            _pbdr(ppr, ('bottom',), BORDER, sz=4, space=4)
        st.paragraph_format.keep_with_next = True

    for lvl, spec in HEADING_SPECS.items():
        _hstyle(_ensure_style(doc, f'Heading {lvl}', WD_STYLE_TYPE.PARAGRAPH,
                              'Normal', builtin=True), *spec)
    try:
        _hstyle(doc.styles['Title'], *HEADING_SPECS[1])
    except KeyError:
        pass

    hl = _ensure_style(doc, 'Hyperlink', WD_STYLE_TYPE.CHARACTER)
    hl.font.color.rgb = ACCENT
    hl.font.underline = False

    q = _ensure_style(doc, 'Quote', WD_STYLE_TYPE.PARAGRAPH, 'Normal', builtin=True)
    q.font.color.rgb = GREY
    q.font.italic = False
    q.font.size = Pt(BODY_PT)
    _rfonts(q.element.get_or_add_rPr(), FONT_LATIN, cjk)
    q.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    q.paragraph_format.left_indent = Pt(14)
    q.paragraph_format.right_indent = Pt(0)
    qppr = q.element.get_or_add_pPr()
    _spacing(qppr, after=after_tw, line=line_tw)
    _pbdr(qppr, ('left',), BORDER, sz=24, space=10)

    cb = _ensure_style(doc, 'GH Code Block', WD_STYLE_TYPE.PARAGRAPH, 'Normal')
    cb.font.name = FONT_MONO
    cb.font.size = Pt(CODE_PT)
    cb.font.color.rgb = INK
    _rfonts(cb.element.get_or_add_rPr(), FONT_MONO, FONT_MONO)
    cbp = cb.element.get_or_add_pPr()
    _spacing(cbp, after='0', line='300')
    _shd(cbp, BG)
    cb.paragraph_format.left_indent = Pt(10)

    ic = _ensure_style(doc, 'GH Inline Code', WD_STYLE_TYPE.CHARACTER)
    ic.font.name = FONT_MONO
    ic.font.size = Pt(CODE_PT)
    icr = ic.element.get_or_add_rPr()
    _rfonts(icr, FONT_MONO, FONT_MONO)
    _shd(icr, CODE_BG)

    for name in ('List Bullet', 'List Bullet 2', 'List Number', 'List Number 2'):
        try:
            ls = doc.styles[name]
        except KeyError:
            continue
        _rfonts(ls.element.get_or_add_rPr(), FONT_LATIN, cjk)
        _spacing(ls.element.get_or_add_pPr(), after='60')


def _table_cell_margins(t):
    tblPr = t._tbl.tblPr
    old = tblPr.find(qn('w:tblCellMar'))
    if old is not None:
        tblPr.remove(old)
    mar = OxmlElement('w:tblCellMar')
    for edge, val in (('top', 90), ('bottom', 90), ('left', 195), ('right', 195)):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    tblPr.append(mar)


def _table_borders(t):
    tblPr = t._tbl.tblPr
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tb = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), BORDER)
        tb.append(e)
    tblPr.append(tb)


def style_table(t, bold_first_column=False, line_tw='276'):
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_cell_margins(t)
    _table_borders(t)
    tblPr = t._tbl.tblPr
    for tag in ('w:tblStyle', 'w:tblLook'):
        el = tblPr.find(qn(tag))
        if el is not None:
            tblPr.remove(el)
    for ri, tr in enumerate(t._tbl.tr_lst):
        fill = 'FFFFFF' if ri == 0 or ri % 2 == 1 else BG
        for tc in tr.tc_lst:
            _shd(tc.get_or_add_tcPr(), fill)
            for pel in tc.findall(qn('w:p')):
                ppr = pel.find(qn('w:pPr'))
                if ppr is None:
                    ppr = OxmlElement('w:pPr')
                    pel.insert(0, ppr)
                _spacing(ppr, before='0', after='0', line=line_tw)
            emph = ri == 0 or bold_first_column and tc is tr.tc_lst[0]
            for r in tc.iter(qn('w:r')):
                rpr = r.find(qn('w:rPr'))
                if rpr is None:
                    rpr = OxmlElement('w:rPr')
                    r.insert(0, rpr)
                for tag in ('w:sz', 'w:szCs'):
                    el = rpr.find(qn(tag))
                    if el is None:
                        el = OxmlElement(tag)
                        rpr.append(el)
                    el.set(qn('w:val'), str(int(TABLE_PT * 2)))
                if emph and rpr.find(qn('w:b')) is None:
                    rpr.append(OxmlElement('w:b'))


def add_footer(doc, text=FOOTER_TEXT, preserve_existing=False):
    for sec in doc.sections:
        if sec.footer.is_linked_to_previous:
            sec.footer.is_linked_to_previous = False
        fp = (sec.footer.paragraphs[0] if sec.footer.paragraphs
              else sec.footer.add_paragraph())
        if preserve_existing and fp.text.strip():
            continue
        fp.text = text
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in fp.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = GREY
            _rfonts(r._element.get_or_add_rPr(), FONT_LATIN, FONT_CJK)


def add_hyperlink(p, url, text):
    rid = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement('w:hyperlink')
    h.set(qn('r:id'), rid)
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rs = OxmlElement('w:rStyle')
    rs.set(qn('w:val'), 'Hyperlink')
    rpr.append(rs)
    r.append(rpr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    h.append(r)
    p._p.append(h)


_TOKEN = re.compile(
    r'\*\*\*(?P<bi>.+?)\*\*\*|\*\*(?P<b>.+?)\*\*|(?<!\*)\*(?P<i>[^*]+)\*'
    r'|`(?P<c>[^`]+)`|~~(?P<s>.+?)~~|\[(?P<lt>[^\]]+)\]\((?P<lu>[^)\s]+)\)')


class GithubDoc:
    def __init__(self, cjk=FONT_CJK, line_mult=1.15, compact=False):
        self.doc = Document()
        apply_github_styles(self.doc, cjk, line_mult, compact)

    def _inline(self, p, text):
        pos = 0
        for m in _TOKEN.finditer(text):
            if m.start() > pos:
                p.add_run(text[pos:m.start()])
            g = m.groupdict()
            if g['bi']:
                r = p.add_run(g['bi']); r.bold = r.italic = True
            elif g['b']:
                p.add_run(g['b']).bold = True
            elif g['i']:
                p.add_run(g['i']).italic = True
            elif g['c']:
                p.add_run(g['c']).style = self.doc.styles['GH Inline Code']
            elif g['s']:
                p.add_run(g['s']).font.strike = True
            elif g['lt']:
                add_hyperlink(p, g['lu'], g['lt'])
            pos = m.end()
        if pos < len(text):
            p.add_run(text[pos:])

    def heading(self, text, level=1):
        p = self.doc.add_paragraph(style=f'Heading {min(max(level, 1), 6)}')
        self._inline(p, text)
        return p

    def para(self, text='', style=None):
        p = self.doc.add_paragraph(style=style)
        if text:
            self._inline(p, text)
        return p

    def quote(self, text):
        return self.para(text, style='Quote')

    def bullets(self, items, ordered=False, level=1):
        base = 'List Number' if ordered else 'List Bullet'
        style = base if level <= 1 else f'{base} {min(level, 2)}'
        for it in items:
            self.para(it, style=style)

    def code_block(self, code):
        lines = code.rstrip('\n').split('\n')
        for i, line in enumerate(lines):
            p = self.doc.add_paragraph(style='GH Code Block')
            p.add_run(line if line else ' ')
            if i == len(lines) - 1:
                _spacing(p._p.get_or_add_pPr(), after=AFTER_TW, line='300')

    def callout(self, kind, text):
        color, hexc, label, _ = CALLOUTS[kind.upper()]
        tp = self.doc.add_paragraph()
        tr = tp.add_run(label)
        tr.bold = True
        tr.font.color.rgb = color
        _spacing(tp._p.get_or_add_pPr(), after='40')
        _pbdr(tp._p.get_or_add_pPr(), ('left',), hexc, sz=24, space=10)
        tp.paragraph_format.left_indent = Pt(14)
        bp = self.para(text)
        _pbdr(bp._p.get_or_add_pPr(), ('left',), hexc, sz=24, space=10)
        bp.paragraph_format.left_indent = Pt(14)
        return bp

    def hr(self):
        p = self.doc.add_paragraph()
        _spacing(p._p.get_or_add_pPr(), before='120', after='120', line=None)
        _pbdr(p._p.get_or_add_pPr(), ('bottom',), BORDER, sz=16, space=1)

    def table(self, rows, bold_first_column=False):
        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        for ri, row in enumerate(rows):
            for ci, v in enumerate(row):
                self._inline(t.cell(ri, ci).paragraphs[0], str(v))
        style_table(t, bold_first_column=bold_first_column)
        sp = self.doc.add_paragraph()
        _spacing(sp._p.get_or_add_pPr(), after='0', line=None)
        return t

    def save(self, path, footer=True):
        if footer:
            add_footer(self.doc)
        self.doc.save(path)
        return path


_H_RE = re.compile(r'^(#{1,6})\s+(.*)$')
_HR_RE = re.compile(r'^(\*{3,}|-{3,}|_{3,})\s*$')
_UL_RE = re.compile(r'^(\s*)[-*+]\s+(.*)$')
_OL_RE = re.compile(r'^(\s*)\d+[.)]\s+(.*)$')
_ALERT_RE = re.compile(r'^\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]\s*$', re.I)


def md_to_doc(gd, md):
    lines = md.replace('\r\n', '\n').split('\n')
    i, n = 0, len(lines)
    while i < n:
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith('```'):
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1
            gd.code_block('\n'.join(buf))
            continue
        m = _H_RE.match(stripped)
        if m:
            gd.heading(m.group(2).strip(), level=len(m.group(1)))
            i += 1
            continue
        if _HR_RE.match(stripped):
            gd.hr()
            i += 1
            continue
        if stripped.startswith('>'):
            buf = []
            while i < n and lines[i].strip().startswith('>'):
                buf.append(lines[i].strip().lstrip('>').strip())
                i += 1
            if buf and _ALERT_RE.match(buf[0]):
                gd.callout(_ALERT_RE.match(buf[0]).group(1),
                           ' '.join(b for b in buf[1:] if b))
            else:
                for b in buf:
                    if b:
                        gd.quote(b)
            continue
        if '|' in stripped and i + 1 < n and re.match(
                r'^\s*\|?[\s:|-]+\|[\s:|-]*$', lines[i + 1]):
            cells = lambda s: [c.strip() for c in s.strip().strip('|').split('|')]
            rows = [cells(stripped)]
            i += 2
            while i < n and '|' in lines[i] and lines[i].strip():
                rows.append(cells(lines[i]))
                i += 1
            w = len(rows[0])
            rows = [r + [''] * (w - len(r)) if len(r) < w else r[:w] for r in rows]
            gd.table(rows)
            continue
        if _UL_RE.match(lines[i]) or _OL_RE.match(lines[i]):
            while i < n:
                mu, mo = _UL_RE.match(lines[i]), _OL_RE.match(lines[i])
                if not (mu or mo):
                    break
                m2 = mu or mo
                gd.bullets([m2.group(2)], ordered=bool(mo),
                           level=2 if len(m2.group(1)) >= 2 else 1)
                i += 1
            continue
        gd.para(stripped)
        i += 1


_STRIP_TAGS = ('w:rFonts', 'w:color', 'w:sz', 'w:szCs', 'w:highlight', 'w:shd')


def _scrub_runs(root, heading_pars):
    for r in root.iter(qn('w:r')):
        rpr = r.find(qn('w:rPr'))
        if rpr is None:
            continue
        tags = list(_STRIP_TAGS)
        pp = r.getparent()
        while pp is not None and pp.tag != qn('w:p'):
            pp = pp.getparent()
        if pp is not None and id(pp) in heading_pars:
            tags += ('w:b', 'w:bCs', 'w:i', 'w:iCs')
        for tag in tags:
            el = rpr.find(qn(tag))
            if el is not None:
                rpr.remove(el)


def _fix_hyperlinks(doc):
    for h in doc.element.body.iter(qn('w:hyperlink')):
        for r in h.iter(qn('w:r')):
            rpr = r.find(qn('w:rPr'))
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                r.insert(0, rpr)
            if rpr.find(qn('w:rStyle')) is None:
                rs = OxmlElement('w:rStyle')
                rs.set(qn('w:val'), 'Hyperlink')
                rpr.insert(0, rs)
            u = rpr.find(qn('w:u'))
            if u is not None:
                rpr.remove(u)



_HEAD_NUM = [
    (re.compile(r'^第[一二三四五六七八九十百\d]+[章篇部]'), 1),
    (re.compile(r'^第[一二三四五六七八九十百\d]+[节条]'), 2),
    (re.compile(r'^[一二三四五六七八九十]+、'), 2),
    (re.compile(r'^[（(][一二三四五六七八九十]+[）)]'), 3),
    (re.compile(r'^\d+(\.\d+)+\s'), 3),
    (re.compile(r'^\d+[、.．]\s*\S'), 2),
]
_HEAD_END_BAD = tuple('。．！？!?；;，,')


def _detect_headings(doc):
    body = doc.styles['Normal'].font.size
    body_pt = body.pt if body else 10.5
    weights = {}
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.size and r.text.strip():
                key = r.font.size.pt
                weights[key] = weights.get(key, 0) + len(r.text)
    if weights:
        body_pt = max(weights, key=weights.get)

    cands = []
    for p in doc.paragraphs:
        name = p.style.name if p.style else ''
        if name.startswith(('Heading', 'Title', 'List', 'Quote', 'GH Code')):
            continue
        text = p.text.strip()
        if not text or len(text) > 60 or text.endswith(_HEAD_END_BAD):
            continue
        if any(pat.match(text) for pat, _ in COLOR_MARKERS):
            continue
        runs = [r for r in p.runs if r.text.strip()]
        if not runs:
            continue
        sizes = [r.font.size.pt for r in runs if r.font.size]
        size = max(sizes) if sizes else body_pt
        all_bold = all(r.font.bold for r in runs)
        if size >= body_pt * 1.12:
            cands.append((p, size, False, text))
        elif all_bold and len(text) <= 40:
            cands.append((p, size, True, text))

    tiers = sorted({size for _, size, bold_only, _ in cands if not bold_only},
                   reverse=True)
    tier_lvl = {s: min(i + 1, 3) for i, s in enumerate(tiers)}
    promoted = set()
    for p, size, bold_only, text in cands:
        if bold_only:
            lvl = min(len(tiers) + 1, 4)
            for pat, nlvl in _HEAD_NUM:
                if pat.match(text):
                    lvl = nlvl
                    break
        else:
            lvl = tier_lvl[size]
        p.style = doc.styles[f'Heading {lvl}']
        promoted.add(id(p._p))
    return promoted

def _colorize(doc):
    for p in doc.paragraphs:
        name = p.style.name if p.style else ''
        if name.startswith(('Heading', 'Title', 'GH Code')) or not p.text.strip():
            continue
        for pat, kind in COLOR_MARKERS:
            m = pat.match(p.text)
            if not m:
                continue
            color, hexc, _, subtle = CALLOUTS[kind]
            ppr = p._p.get_or_add_pPr()
            _pbdr(ppr, ('left',), hexc, sz=24, space=10)
            _shd(ppr, subtle)
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.right_indent = Pt(8)
            if p.runs and p.runs[0].text == m.group(0):
                p.runs[0].bold = True
                p.runs[0].font.color.rgb = color
            break


def restyle_docx(src, out, cjk=FONT_CJK, bold_first_column=False,
                 preserve_footer=False, line_mult=1.15, compact=False,
                 colorize=True, detect_headings=True):
    src_path = Path(src)
    if not src_path.exists():
        raise FileNotFoundError(f'源文件不存在: {src_path}')
    if src_path.suffix.lower() != '.docx':
        raise ValueError(f'不是 .docx 文件: {src_path}')

    doc = Document(src_path)
    promoted = _detect_headings(doc) if detect_headings else set()
    apply_github_styles(doc, cjk, line_mult, compact)

    heading_ids = promoted | {id(p._p) for p in doc.paragraphs
                              if p.style and p.style.name.startswith('Heading')}
    _scrub_runs(doc.element.body, heading_ids)
    for sec in doc.sections:
        _scrub_runs(sec.header._element, set())
        if preserve_footer:
            _scrub_runs(sec.footer._element, set())

    for p in doc.paragraphs:
        ppr = p._p.find(qn('w:pPr'))
        if ppr is None:
            continue
        name = p.style.name if p.style else ''
        if name.startswith(('Heading', 'Quote', 'GH Code')):
            continue
        for tag in ('w:spacing', 'w:shd'):
            el = ppr.find(qn(tag))
            if el is not None:
                ppr.remove(el)

    _fix_hyperlinks(doc)
    line_tw = str(int(round(line_mult * 240)))
    for t in doc.tables:
        style_table(t, bold_first_column=bold_first_column, line_tw=line_tw)
    if colorize:
        _colorize(doc)
    add_footer(doc, preserve_existing=preserve_footer)
    doc.save(out)


def main(argv=None):
    parser = argparse.ArgumentParser(prog='github_light_word_engine')
    parser.add_argument('--cjk-font', default=FONT_CJK)
    parser.add_argument('--line-spacing', type=float, default=1.15)
    parser.add_argument('--compact', action='store_true')
    sub = parser.add_subparsers(dest='cmd', required=True)
    g = sub.add_parser('gen')
    g.add_argument('input')
    g.add_argument('-o', '--output')
    r = sub.add_parser('restyle')
    r.add_argument('input')
    r.add_argument('-o', '--output')
    r.add_argument('--bold-first-col', action='store_true')
    r.add_argument('--preserve-footer', action='store_true')
    r.add_argument('--no-colorize', action='store_true')
    r.add_argument('--no-detect-headings', action='store_true')
    args = parser.parse_args(argv)
    try:
        src = Path(args.input)
        if args.cmd == 'gen':
            out = args.output or src.with_suffix('.docx')
            gd = GithubDoc(args.cjk_font, args.line_spacing, args.compact)
            md_to_doc(gd, src.read_text(encoding='utf-8'))
            gd.save(out)
        else:
            out = args.output or src.parent / f'{src.stem}_github{src.suffix}'
            restyle_docx(src, out, cjk=args.cjk_font,
                         bold_first_column=args.bold_first_col,
                         preserve_footer=args.preserve_footer,
                         line_mult=args.line_spacing, compact=args.compact,
                         colorize=not args.no_colorize,
                         detect_headings=not args.no_detect_headings)
        print(f'✓ 已生成: {out}')
        return 0
    except (FileNotFoundError, ValueError) as e:
        print(f'✗ 错误: {e}', file=sys.stderr)
        return 1


if __name__ == '__main__':
    sys.exit(main())
